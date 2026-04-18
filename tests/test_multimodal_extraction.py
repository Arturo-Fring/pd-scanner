"""Regression coverage for multimodal extraction and detector composition."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

import fitz
from docx import Document
from PIL import Image

from pd_scanner.core.config import AppConfig
from pd_scanner.core.models import ExtractedChunk, RawFinding
from pd_scanner.detectors.base import BaseDetector
from pd_scanner.detectors.detection_pipeline import DetectionPipeline
from pd_scanner.extractors.docx_extractor import DOCXExtractor
from pd_scanner.extractors.html_extractor import HTMLExtractor
from pd_scanner.extractors.ocr_service import OCRResult
from pd_scanner.extractors.pdf_extractor import PDFExtractor


def build_config(tmp_path: Path, mode: str = "deep") -> AppConfig:
    config = AppConfig.build(
        input_path=tmp_path,
        output_path=tmp_path / "out",
        mode=mode,
        workers=1,
    )
    config.ocr.min_pdf_text_chars = 5
    return config


def _mock_ocr(monkeypatch, text: str = "OCR_TEXT") -> None:
    monkeypatch.setattr("pd_scanner.extractors.ocr_service.OCRService.get_status", lambda self: (True, "mocked OCR"))
    monkeypatch.setattr(
        "pd_scanner.extractors.ocr_service.OCRService.extract_text_from_image",
        lambda self, image, lang=None: OCRResult(
            text=text,
            available=True,
            backend="mock_ocr",
            warnings=[],
            metadata={"backend": "mock_ocr", "lang": lang or self.config.ocr.lang},
        ),
    )
    monkeypatch.setattr("pd_scanner.extractors.ocr_service.OCRService.extract_text", lambda self, image: text)


def test_pdf_image_ocr(tmp_path: Path, monkeypatch) -> None:
    _mock_ocr(monkeypatch, text="passport 1234 567890")
    image_path = tmp_path / "embedded.png"
    Image.new("RGB", (120, 60), color="white").save(image_path)

    pdf_path = tmp_path / "image_only.pdf"
    document = fitz.open()
    page = document.new_page()
    page.insert_image(fitz.Rect(50, 50, 170, 110), filename=str(image_path))
    document.save(pdf_path)
    document.close()

    extractor = PDFExtractor(build_config(tmp_path, mode="deep"))
    result = extractor.extract(pdf_path)
    source_types = {chunk.source_type for chunk in result.extracted_text_chunks}

    assert "pdf_image_ocr" in source_types
    assert result.metadata["embedded_images"] >= 1


def test_docx_image_ocr(tmp_path: Path, monkeypatch) -> None:
    _mock_ocr(monkeypatch, text="snils 112-233-445 95")
    image_stream = BytesIO()
    Image.new("RGB", (120, 60), color="white").save(image_stream, format="PNG")
    image_stream.seek(0)

    docx_path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("Имя: Иван Петров")
    document.add_picture(image_stream)
    document.save(docx_path)

    extractor = DOCXExtractor(build_config(tmp_path, mode="deep"))
    result = extractor.extract(docx_path)
    source_types = {chunk.source_type for chunk in result.extracted_text_chunks}

    assert "docx_image_ocr" in source_types
    assert "docx_paragraph" in source_types
    assert result.metadata["embedded_images"] >= 1


def test_html_image_ocr(tmp_path: Path, monkeypatch) -> None:
    _mock_ocr(monkeypatch, text="inn 500100732259")
    image_path = tmp_path / "page_image.png"
    Image.new("RGB", (100, 40), color="white").save(image_path)
    html_path = tmp_path / "sample.html"
    html_path.write_text(
        """
        <html>
          <head>
            <meta name="description" content="employee directory">
          </head>
          <body>
            <a href="/users/1" title="profile link">Open profile</a>
            <img src="page_image.png" alt="passport scan" title="embedded card">
            <p>Email: ivan.petrov@mail.ru</p>
          </body>
        </html>
        """,
        encoding="utf-8",
    )

    extractor = HTMLExtractor(build_config(tmp_path, mode="deep"))
    result = extractor.extract(html_path)
    source_types = {chunk.source_type for chunk in result.extracted_text_chunks}

    assert "html_text" in source_types
    assert "html_alt_text" in source_types
    assert "html_link" in source_types
    assert "html_metadata" in source_types
    assert "html_image_ocr" in source_types


def test_source_type_assignment(tmp_path: Path, monkeypatch) -> None:
    _mock_ocr(monkeypatch, text="ocr sample")
    image_stream = BytesIO()
    Image.new("RGB", (120, 60), color="white").save(image_stream, format="PNG")
    image_stream.seek(0)

    docx_path = tmp_path / "typed.docx"
    document = Document()
    document.add_paragraph("Контакт: +7 999 123-45-67")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "email"
    document.add_picture(image_stream)
    document.save(docx_path)

    extractor = DOCXExtractor(build_config(tmp_path, mode="deep"))
    result = extractor.extract(docx_path)
    source_types = {chunk.source_type for chunk in result.extracted_text_chunks}

    assert {"docx_paragraph", "docx_table_cell", "docx_image_ocr"} <= source_types


class _StaticDetector(BaseDetector):
    def __init__(self, findings: list[RawFinding]) -> None:
        self._findings = findings

    def detect(self, chunks: list[ExtractedChunk]) -> list[RawFinding]:
        _ = chunks
        return list(self._findings)


def test_detector_pipeline_merges_duplicate_span() -> None:
    chunk = ExtractedChunk(text="Email: ivan.petrov@mail.ru", source_type="html_text", source_path="sample.html")
    findings_a = [
        RawFinding(
            entity_type="email",
            group="common",
            original_value="ivan.petrov@mail.ru",
            normalized_value="ivan.petrov@mail.ru",
            masked_value="iv***@mail.ru",
            confidence=0.72,
            explanation="rule hit",
            source_context="Email: iv***@mail.ru",
            row_key='{"page": 1}',
            start=7,
            end=26,
            source_detector="rule_based",
            chunk_source_type="html_text",
            source_path="sample.html",
        )
    ]
    findings_b = [
        RawFinding(
            entity_type="email",
            group="common",
            original_value="ivan.petrov@mail.ru",
            normalized_value="ivan.petrov@mail.ru",
            masked_value="iv***@mail.ru",
            confidence=0.91,
            explanation="model hit",
            source_context="Email: iv***@mail.ru",
            row_key='{"page": 1}',
            start=7,
            end=26,
            source_detector="model_stub",
            chunk_source_type="html_text",
            source_path="sample.html",
        )
    ]

    pipeline = DetectionPipeline([_StaticDetector(findings_a), _StaticDetector(findings_b)])
    merged = pipeline.detect([chunk])

    assert len(merged) == 1
    assert merged[0].confidence == 0.91
    assert merged[0].source_detector == "model_stub,rule_based"
    assert "model hit" in merged[0].explanation
