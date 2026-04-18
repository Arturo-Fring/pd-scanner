"""DOCX extractor."""

from __future__ import annotations

import io
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from PIL import Image

from pd_scanner.core.models import ExtractedChunk, ExtractionResult
from pd_scanner.core.utils import sanitize_whitespace
from pd_scanner.extractors.base import BaseExtractor
from pd_scanner.extractors.resource_router import EmbeddedResource


class DOCXExtractor(BaseExtractor):
    """Extract paragraphs and tables from DOCX files."""

    file_type = "docx"

    def extract(self, path: Path) -> ExtractionResult:
        document = Document(path)
        chunks: list[ExtractedChunk] = []
        warnings: list[str] = []
        ocr_attempts = 0
        ocr_successes = 0
        image_count = 0
        use_ocr = self.config.runtime.mode == "deep"
        ocr_status_payload = self.ocr_service.get_status_payload() if use_ocr else None
        ocr_available = bool(ocr_status_payload["available"]) if ocr_status_payload else False
        ocr_status = str(ocr_status_payload["message"]) if ocr_status_payload else "OCR disabled in fast mode."
        for index, paragraph in enumerate(document.paragraphs):
            text = sanitize_whitespace(paragraph.text)
            if text:
                routed_chunks, routed_warnings = self.route_resource(
                    EmbeddedResource(
                        resource_type="text",
                        payload=text,
                        source_type="docx_paragraph",
                        source_path=str(path),
                        location={"paragraph": index + 1},
                        metadata={"paragraph": index + 1},
                    )
                )
                chunks.extend(routed_chunks)
                warnings.extend(routed_warnings)
        for table_index, table in enumerate(document.tables):
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    cell_text = sanitize_whitespace(cell.text)
                    if not cell_text:
                        continue
                    routed_chunks, routed_warnings = self.route_resource(
                        EmbeddedResource(
                            resource_type="text",
                            payload=cell_text,
                            source_type="docx_table_cell",
                            source_path=str(path),
                            location={"table": table_index + 1, "row": row_index + 1, "cell": cell_index + 1},
                            metadata={"structured": True, "table": table_index + 1, "row": row_index + 1, "cell": cell_index + 1},
                        )
                    )
                    for chunk in routed_chunks:
                        chunk.columns = ("docx_table_cell",)
                    chunks.extend(routed_chunks)
                    warnings.extend(routed_warnings)

        for image_index, relationship in enumerate(document.part._rels.values(), start=1):
            if relationship.reltype != RT.IMAGE:
                continue
            image_count += 1
            if image_count > self.config.runtime.max_embedded_images_per_file:
                warnings.append("DOCX embedded image limit reached; remaining images skipped.")
                break
            if ocr_attempts >= self.config.runtime.max_ocr_calls_per_file:
                warnings.append("DOCX OCR call limit reached; remaining images skipped.")
                break
            if not use_ocr:
                continue
            if not ocr_available:
                warnings.append(f"DOCX embedded image OCR skipped: {ocr_status}")
                break
            try:
                with Image.open(io.BytesIO(relationship.target_part.blob)) as image:
                    route_result = self.route_resource_detailed(
                        EmbeddedResource(
                            resource_type="image",
                            payload=image.copy(),
                            source_type="docx_image_ocr",
                            source_path=str(path),
                            location={"image_index": image_index},
                            metadata={"image_index": image_index},
                        )
                    )
                warnings.extend(route_result.warnings)
                if route_result.attempted_ocr:
                    ocr_attempts += 1
                if route_result.chunks:
                    chunks.extend(route_result.chunks)
                if route_result.ocr_text_found:
                    ocr_successes += 1
                if any(self.ocr_service.is_runtime_failure_warning(item) for item in route_result.warnings):
                    warnings.append("DOCX embedded image OCR disabled for remaining images after backend issue.")
                    break
            except Exception as exc:
                warnings.append(f"DOCX embedded image OCR failed for image {image_index}: {exc}")

        return self.build_result(
            chunks=chunks,
            metadata={
                "chunk_count": len(chunks),
                "paragraph_count": len(document.paragraphs),
                "table_count": len(document.tables),
                "embedded_images": image_count,
                "ocr_available": ocr_available,
                "ocr_status": ocr_status,
                "ocr_backend": ocr_status_payload.get("backend") if ocr_status_payload else None,
                "ocr_device": ocr_status_payload.get("device") if ocr_status_payload else None,
                "ocr_calls": ocr_attempts,
                "ocr_successes": ocr_successes,
                "ocr_used": ocr_attempts > 0,
                "structured": False,
            },
            warnings=warnings,
        )
